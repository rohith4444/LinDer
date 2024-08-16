import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
from datetime import datetime
import wave
from pydub import AudioSegment
from logging_config import get_module_logger
from config.settings import (
    LANGUAGES, OPENAI_API_KEY, GOOGLE_APPLICATION_CREDENTIALS,
    AUDIO_OUTPUT_DIR, DOCUMENT_INPUT_DIR, DOCUMENT_OUTPUT_DIR
)

# Get logger for this module
logger = get_module_logger(__name__)

def generate_unique_filename(base_name, extension):
    """
    Generates a unique filename using the current timestamp.
    
    :param base_name: The base name for the file
    :param extension: The file extension
    :return: A unique filename
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{base_name}_{timestamp}{extension}"
    logger.info(f"Generated unique filename: {unique_filename}")
    return unique_filename

def get_language_choice(prompt, languages=LANGUAGES):
    """
    Presents a list of language choices to the user and returns the selected language.
    
    :param prompt: The prompt to display to the user
    :param languages: A dictionary of language choices (default: LANGUAGES from settings)
    :return: The selected language
    """
    logger.info("Presenting language choices to user")
    print(prompt)
    for key, language in languages.items():
        print(f"{key}. {language[0]}")
    
    while True:
        choice = input("Enter your choice: ")
        if choice in languages:
            selected_language = languages[choice]
            logger.info(f"User selected language: {selected_language[0]}")
            return selected_language
        logger.warning(f"Invalid language choice: {choice}")
        print("Invalid choice. Please try again.")

def get_filename(default_name, extension):
    """
    Prompts the user for a filename and returns it with the given extension.
    
    :param default_name: The default filename to use if the user doesn't provide one
    :param extension: The file extension to use
    :return: The filename with extension
    """
    filename = input(f"Enter a filename (default: {default_name}{extension}): ").strip()
    if not filename:
        filename = default_name
    if not filename.endswith(extension):
        filename += extension
    logger.info(f"User provided filename: {filename}")
    return filename

def load_env_variables():
    """
    Loads environment variables from a .env file.
    """
    logger.info("Loading environment variables")
    load_dotenv()

    # Check for required environment variables
    required_vars = {
        "OPENAI_API_KEY": OPENAI_API_KEY,
        "GOOGLE_APPLICATION_CREDENTIALS": GOOGLE_APPLICATION_CREDENTIALS
    }

    for var, value in required_vars.items():
        if not value:
            logger.error(f"Missing required environment variable: {var}")
            raise EnvironmentError(f"Missing required environment variable: {var}")
    
    logger.info("Environment variables loaded successfully")

def read_text_file(file_path):
    """
    Reads content from a text file.
    
    :param file_path: Path to the text file
    :return: Content of the text file
    """
    logger.info(f"Reading text file: {file_path}")
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        logger.info("Text file read successfully")
        return content
    except Exception as e:
        logger.exception(f"Error reading text file: {str(e)}")
        raise

def read_pdf_file(file_path):
    """
    Reads content from a PDF file.
    
    :param file_path: Path to the PDF file
    :return: Content of the PDF file
    """
    logger.info(f"Reading PDF file: {file_path}")
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        logger.info("PDF file read successfully")
        return text
    except Exception as e:
        logger.exception(f"Error reading PDF file: {str(e)}")
        raise

def read_docx_file(file_path):
    """
    Reads content from a DOCX file.
    
    :param file_path: Path to the DOCX file
    :return: Content of the DOCX file
    """
    logger.info(f"Reading DOCX file: {file_path}")
    try:
        doc = Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        logger.info("DOCX file read successfully")
        return content
    except Exception as e:
        logger.exception(f"Error reading DOCX file: {str(e)}")
        raise

def split_content(content, max_chars=5000):
    """
    Splits content into chunks of maximum characters, trying to break at sentence boundaries.
    
    :param content: The content to split
    :param max_chars: Maximum characters per chunk (default: 5000)
    :return: List of content chunks
    """
    logger.info(f"Splitting content into chunks of max {max_chars} characters")
    sentences = content.replace('\n', ' ').split('. ')
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 2 <= max_chars:  # +2 for the '. ' we removed
            current_chunk += sentence + '. '
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence + '. '

    if current_chunk:
        chunks.append(current_chunk.strip())

    logger.info(f"Content split into {len(chunks)} chunks")
    return chunks

def read_large_file(file_path, chunk_size=1024*1024):
    """
    Reads a large file in chunks.
    
    :param file_path: Path to the file
    :param chunk_size: Size of each chunk in bytes (default: 1MB)
    :yield: Chunks of the file content
    """
    logger.info(f"Reading large file in chunks: {file_path}")
    with open(file_path, 'rb') as file:
        while True:
            chunk = file.read(chunk_size)
            if not chunk:
                break
            yield chunk

def read_file(file_path):
    """
    Reads content from a file based on its extension.
    
    :param file_path: Path to the file
    :return: Content of the file
    """
    logger.info(f"Reading file: {file_path}")
    _, file_extension = os.path.splitext(file_path)
    file_size = os.path.getsize(file_path)
    
    if file_size > 10 * 1024 * 1024:  # If file is larger than 10MB
        logger.info(f"Large file detected ({file_size} bytes). Reading in chunks.")
        return b''.join(read_large_file(file_path))
    
    if file_extension.lower() == '.txt':
        return read_text_file(file_path)
    elif file_extension.lower() == '.pdf':
        return read_pdf_file(file_path)
    elif file_extension.lower() == '.docx':
        return read_docx_file(file_path)
    else:
        logger.error(f"Unsupported file format: {file_extension}")
        raise ValueError(f"Unsupported file format: {file_extension}")

def write_pdf(content, output_file):
    """
    Writes content to a PDF file.
    
    :param content: The content to write to the PDF
    :param output_file: The path to save the PDF file
    """
    logger.info(f"Writing content to PDF file: {output_file}")
    try:
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        width, height = letter
        y = height - 50  # Start 50 points down from the top

        for line in content.split('\n'):
            if y < 50:  # If we're near the bottom of the page
                can.showPage()  # Start a new page
                y = height - 50  # Reset y to the top of the new page

            can.drawString(50, y, line)
            y -= 15  # Move down 15 points

        can.save()

        packet.seek(0)
        new_pdf = PdfReader(packet)
        writer = PdfWriter()
        writer.add_page(new_pdf.pages[0])

        with open(output_file, 'wb') as output_file_handle:
            writer.write(output_file_handle)
        logger.info("Content written to PDF successfully")
    except Exception as e:
        logger.exception(f"Error writing to PDF file: {str(e)}")
        raise

def write_docx(content, output_file):
    """
    Writes content to a DOCX file.
    
    :param content: The content to write to the DOCX
    :param output_file: The path to save the DOCX file
    """
    logger.info(f"Writing content to DOCX file: {output_file}")
    try:
        doc = Document()
        doc.add_paragraph(content)
        doc.save(output_file)
        logger.info("Content written to DOCX successfully")
    except Exception as e:
        logger.exception(f"Error writing to DOCX file: {str(e)}")
        raise

def write_large_file(content, output_file, chunk_size=1024*1024):
    """
    Writes large content to a file in chunks.
    
    :param content: The content to write
    :param output_file: The path to save the file
    :param chunk_size: Size of each chunk in bytes (default: 1MB)
    """
    logger.info(f"Writing large file in chunks: {output_file}")
    with open(output_file, 'wb') as file:
        for i in range(0, len(content), chunk_size):
            file.write(content[i:i+chunk_size])

def write_file(content, output_file):
    """
    Writes content to a file based on its extension.
    
    :param content: The content to write to the file
    :param output_file: The path to save the file
    """
    logger.info(f"Writing content to file: {output_file}")
    _, file_extension = os.path.splitext(output_file)
    try:
        if isinstance(content, str) and len(content.encode('utf-8')) > 10 * 1024 * 1024:
            write_large_file(content.encode('utf-8'), output_file)
        elif isinstance(content, bytes) and len(content) > 10 * 1024 * 1024:
            write_large_file(content, output_file)
        elif file_extension.lower() == '.txt':
            with open(output_file, 'w', encoding='utf-8') as file:
                file.write(content)
            logger.info("Content written to text file successfully")
        elif file_extension.lower() == '.pdf':
            write_pdf(content, output_file)
        elif file_extension.lower() == '.docx':
            write_docx(content, output_file)
        else:
            logger.error(f"Unsupported output file format: {file_extension}")
            raise ValueError(f"Unsupported output file format: {file_extension}")
    except Exception as e:
        logger.exception(f"Error writing to file: {str(e)}")
        raise

def check_text_size(text):
    """
    Checks the size of the input text.
    
    :param text: The input text (string or bytes)
    :return: True if the text is considered large, False otherwise
    """
    logger.info("Checking text size")
    try:
        if isinstance(text, str):
            size = len(text.encode('utf-8'))
        elif isinstance(text, bytes):
            size = len(text)
        else:
            raise ValueError("Input must be string or bytes")
        
        # Consider text large if it's more than 4000 characters (OpenAI's typical limit)
        is_large = size > 3500
        logger.info(f"Text size: {size} bytes. Considered large: {is_large}")
        return is_large
    except Exception as e:
        logger.exception(f"An error occurred while checking text size: {str(e)}")
        raise

def check_audio_duration(audio_file):
    """
    Checks the duration of the input audio file.
    
    :param audio_file: Path to the audio file
    :return: True if the audio is considered large, False otherwise
    """
    logger.info(f"Checking audio duration for file: {audio_file}")
    try:
        _, file_extension = os.path.splitext(audio_file)
        
        if file_extension.lower() == '.wav':
            with wave.open(audio_file, 'rb') as wav_file:
                frames = wav_file.getnframes()
                rate = wav_file.getframerate()
                duration = frames / float(rate)
        elif file_extension.lower() in ['.mp3', '.ogg', '.flac']:
            audio = AudioSegment.from_file(audio_file)
            duration = len(audio) / 1000.0  # pydub works in milliseconds
        else:
            raise ValueError(f"Unsupported audio format: {file_extension}")
        
        # Consider audio large if it's longer than 1 minute
        is_large = duration > 60
        logger.info(f"Audio duration: {duration:.2f} seconds. Considered large: {is_large}")
        return is_large
    except Exception as e:
        logger.exception(f"An error occurred while checking audio duration: {str(e)}")
        raise