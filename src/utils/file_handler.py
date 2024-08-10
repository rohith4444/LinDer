import os
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
from text.text_processor import translate_text

def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def read_pdf_file(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def read_docx_file(file_path):
    doc = Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def read_file(file_path):
    _, file_extension = os.path.splitext(file_path)
    if file_extension.lower() == '.txt':
        return read_text_file(file_path)
    elif file_extension.lower() == '.pdf':
        return read_pdf_file(file_path)
    elif file_extension.lower() == '.docx':
        return read_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def write_pdf(content, output_file):
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    width, height = letter
    y = height - 50  # Start 50 points down from the top

    for line in content.split('\n'):
        if y < 50:  # If we're near the bottom of the page
            can.showPage()  # Start a new page
            y = height - 50  # Reset y to the top of the new page

        can.drawString(50, y, line)
        y -= 15  # Move down 15 points

    can.save()

    packet.seek(0)
    new_pdf = PdfReader(packet)
    writer = PdfWriter()
    writer.add_page(new_pdf.pages[0])

    with open(output_file, 'wb') as output_file_handle:
        writer.write(output_file_handle)

def translate_file(input_file, output_file, source_lang, target_lang):
    content = read_file(input_file)
    translated_content = translate_text(content, source_lang, target_lang)
    
    _, file_extension = os.path.splitext(output_file)
    if file_extension.lower() == '.txt':
        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(translated_content)
    elif file_extension.lower() == '.pdf':
        write_pdf(translated_content, output_file)
    elif file_extension.lower() == '.docx':
        doc = Document()
        doc.add_paragraph(translated_content)
        doc.save(output_file)
    else:
        raise ValueError(f"Unsupported output file format: {file_extension}")

    return output_file